#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Word文档纯文本转换工具
功能：将Word文档内容转为纯文本格式，再写回原Word文档
操作等效于：
    1. 打开docx文档
    2. 按ctrl+a全选所有内容
    3. 将内容粘贴到记事本
    4. 再把内容从记事本复制回原Word文档
"""

import os
import docx
from docx.shared import Pt
import glob
import shutil
from tqdm import tqdm
import sys


def process_word_file(input_file, backup=True):
    """
    处理单个Word文件
    
    Args:
        input_file: Word文件路径
        backup: 是否创建备份
    
    Returns:
        成功返回True，失败返回False
    """
    try:
        # 创建备份
        if backup:
            backup_file = input_file + ".bak"
            shutil.copy2(input_file, backup_file)
            print(f"已创建备份: {backup_file}")
        
        # 读取Word文档内容
        doc = docx.Document(input_file)
        
        # 提取纯文本
        text_content = []
        for para in doc.paragraphs:
            text_content.append(para.text)
        
        # 将所有文本合并，保留段落分隔
        full_text = "\n".join(text_content)
        
        # 创建新的空白文档
        new_doc = docx.Document()
        
        # 设置默认字体和字号
        style = new_doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # 将纯文本内容写入新文档，保留段落分隔
        for line in full_text.split('\n'):
            new_doc.add_paragraph(line)
        
        # 保存文档
        new_doc.save(input_file)
        
        return True
    
    except Exception as e:
        print(f"处理文件 {input_file} 时出错: {str(e)}")
        return False


def batch_process_word_files(folder_path, backup=True):
    """
    批量处理文件夹中的所有Word文档
    
    Args:
        folder_path: 文件夹路径
        backup: 是否创建备份
    """
    # 查找所有Word文档
    file_pattern = os.path.join(folder_path, "**", "*.docx")
    word_files = glob.glob(file_pattern, recursive=True)
    
    if not word_files:
        print(f"在 {folder_path} 及其子文件夹中未找到任何Word文档(.docx)。")
        return
    
    print(f"找到 {len(word_files)} 个Word文档(.docx)。")
    
    # 处理每个文件
    success_count = 0
    
    for file_path in tqdm(word_files, desc="处理进度"):
        if process_word_file(file_path, backup):
            success_count += 1
    
    print(f"\n处理完成！成功处理 {success_count}/{len(word_files)} 个文件。")
    
    if success_count < len(word_files):
        print("部分文件处理失败，请查看上方错误信息。")


def main():
    """主函数"""
    print("Word文档纯文本转换工具")
    print("=" * 40)
    
    # 检查命令行参数
    if len(sys.argv) > 1:
        folder_path = sys.argv[1]
    else:
        # 提示用户输入文件夹路径
        folder_path = input("请输入Word文档所在的文件夹路径（默认为当前目录）: ").strip()
        if not folder_path:
            folder_path = os.getcwd()
    
    # 检查文件夹是否存在
    if not os.path.isdir(folder_path):
        print(f"错误: 文件夹 '{folder_path}' 不存在！")
        return
    
    # 询问是否创建备份
    backup_response = input("是否为每个Word文档创建备份？(y/n，默认为y): ").strip().lower()
    create_backup = backup_response != 'n'
    
    if create_backup:
        print("将为每个处理的文档创建.bak备份文件。")
    
    # 确认操作
    confirm = input(f"确认要处理文件夹 '{folder_path}' 中的所有Word文档吗？(y/n): ").strip().lower()
    if confirm != 'y':
        print("操作已取消。")
        return
    
    # 执行批处理
    batch_process_word_files(folder_path, create_backup)


if __name__ == "__main__":
    main() 